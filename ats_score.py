# -*- coding: utf-8 -*-
"""ats score

Automatically generated by Colab.

Original file is located at
    https://colab.research.google.com/drive/1OUunCTmsopbUVLq-EokzNQiRZ5EBRZT6
"""

!pip install nltk spacy pdfplumber docx2txt python-docx
!python -m spacy download en_core_web_sm

import nltk
import os
import re
import spacy
import pdfplumber
import docx2txt
from nltk.tokenize import word_tokenize
from nltk.corpus import stopwords
from docx import Document
from collections import Counter
nltk.download('punkt')
nltk.download('stopwords')
nltk.download('punkt_tab')
print(nltk.data.path)

rm -rf ~/.cache/nltk

nlp = spacy.load('en_core_web_sm')

def extract_text_from_file(file_path):
    """
    Extracts text from PDF, Word, or text files.
    """
    text = ''
    if file_path.endswith('.pdf'):
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + '\n'
    elif file_path.endswith('.docx') or file_path.endswith('.doc'):
        text = docx2txt.process(file_path)
    elif file_path.endswith('.txt'):
        with open(file_path, 'r', encoding='utf-8') as file:
            text = file.read()
    else:
        raise ValueError('Unsupported file format: ' + file_path)
    return text

def preprocess_text(text):
    """
    Tokenizes and cleans text for analysis.
    """
    tokens = word_tokenize(text)
    tokens = [token.lower() for token in tokens if token.isalpha()]
    stop_words = set(stopwords.words('english'))
    tokens = [token for token in tokens if token not in stop_words]
    return tokens

def calculate_ats_score(resume_text, job_desc_text):
    """
    Calculates the ATS score by comparing keywords between the resume and the job description.
    """
    resume_tokens = preprocess_text(resume_text)
    job_desc_tokens = preprocess_text(job_desc_text)

    # Use Counter to get frequency of each token
    resume_counts = Counter(resume_tokens)
    job_desc_counts = Counter(job_desc_tokens)

    # Calculate keyword matches
    matched_keywords = set(resume_tokens) & set(job_desc_tokens)
    total_keywords = len(set(job_desc_tokens))
    if total_keywords == 0:
        keyword_score = 0
    else:
        keyword_score = len(matched_keywords) / total_keywords * 100  # Percentage

    # Simple formatting score based on presence of sections
    formatting_sections = ['Education', 'Experience', 'Skills']
    formatting_score = sum(1 for section in formatting_sections if section.lower() in resume_text.lower())
    formatting_score = formatting_score / len(formatting_sections) * 100  # Percentage

    # Combined ATS score (weighted)
    ats_score = (0.7 * keyword_score) + (0.3 * formatting_score)
    return ats_score, matched_keywords

def provide_suggestions(ats_score, matched_keywords, job_desc_text, resume_text):
    """
    Provides suggestions to improve the resume based on the ATS score.
    """
    suggestions = []
    # Define thresholds
    ats_threshold = 70  # Consider ATS score below 70 as low
    if ats_score < ats_threshold:
        suggestions.append(f"Your ATS score is {ats_score:.2f}%, which is below the optimal threshold of {ats_threshold}%. Consider the following improvements:")
        # Identify missing keywords
        job_desc_tokens = set(preprocess_text(job_desc_text))
        resume_tokens = set(preprocess_text(resume_text))
        missing_keywords = job_desc_tokens - resume_tokens
        if missing_keywords:
            suggestions.append(f"- Add missing keywords relevant to the job description: {', '.join(missing_keywords)}")
        else:
            suggestions.append("- Ensure that your resume highlights relevant experience and skills.")
        # Check for formatting improvements
        formatting_sections = ['Education', 'Experience', 'Skills']
        for section in formatting_sections:
            if section.lower() not in resume_text.lower():
                suggestions.append(f"- Add a '{section}' section to your resume.")
    else:
        suggestions.append(f"Your ATS score is {ats_score:.2f}%, which is good.")
    return suggestions

def rewrite_resume(resume_text, suggestions, job_desc_text):
    """
    Rewrites the resume based on the suggestions provided.
    """
    # Parse missing keywords from suggestions
    missing_keywords = []
    for suggestion in suggestions:
        if 'Add missing keywords' in suggestion:
            parts = suggestion.split(':')
            if len(parts) > 1:
                keywords_str = parts[1].strip()
                missing_keywords = [kw.strip() for kw in keywords_str.split(',')]
                break

    # Add missing keywords to the 'Skills' section (or create it if it doesn't exist)
    resume_sections = split_resume_sections(resume_text)
    if 'Skills' in resume_sections:
        # Append missing keywords
        current_skills = resume_sections['Skills']
        updated_skills = current_skills + '\n' + ', '.join(missing_keywords)
        resume_sections['Skills'] = updated_skills
    else:
        # Create 'Skills' section
        resume_sections['Skills'] = ', '.join(missing_keywords)

    # Add missing sections
    for suggestion in suggestions:
        if suggestion.startswith("- Add a '") and suggestion.endswith("' section to your resume."):
            section_name = suggestion[9:-27]
            resume_sections[section_name] = ""  # Empty section

    # Reconstruct the resume text
    new_resume_text = reconstruct_resume(resume_sections)
    return new_resume_text

def split_resume_sections(resume_text):
    """
    Splits the resume into sections based on headings.
    """
    lines = resume_text.split('\n')
    sections = {}
    current_section = None
    for line in lines:
        line_stripped = line.strip()
        if line_stripped.isupper() and len(line_stripped.split()) <= 3:
            # Heading detected
            current_section = line_stripped
            sections[current_section] = ''
        elif current_section:
            sections[current_section] += line + '\n'
    return sections

def reconstruct_resume(sections):
    """
    Reconstructs the resume text from sections.
    """
    resume_text = ''
    for section, content in sections.items():
        resume_text += section + '\n'
        resume_text += content.strip() + '\n\n'
    return resume_text.strip()

def main():
    # Paths to the candidate resume and job description
    resume_path = 'sp.txt'  # Replace with your resume file path
    job_desc_path = 'jd.txt'  # Replace with your job description file path

    # Extract texts
    resume_text = extract_text_from_file(resume_path)
    job_desc_text = extract_text_from_file(job_desc_path)

    # Calculate initial ATS score
    ats_score, matched_keywords = calculate_ats_score(resume_text, job_desc_text)
    print(f"Initial ATS Score: {ats_score:.2f}%")
    print(f"Matched Keywords: {', '.join(matched_keywords)}")

    # Provide suggestions
    suggestions = provide_suggestions(ats_score, matched_keywords, job_desc_text, resume_text)
    print("\nSuggestions:")
    for suggestion in suggestions:
        print(suggestion)

    # Rewrite resume based on suggestions
    if ats_score < 70:
        new_resume_text = rewrite_resume(resume_text, suggestions, job_desc_text)
        # Recalculate ATS score after adjustments
        new_ats_score, new_matched_keywords = calculate_ats_score(new_resume_text, job_desc_text)
        print(f"\nNew ATS Score: {new_ats_score:.2f}%")
        print(f"New Matched Keywords: {', '.join(new_matched_keywords)}")
        # Save the new resume
        save_resume_to_file(new_resume_text, 'improved_resume.docx')
        print("\nYour improved resume has been saved as 'improved_resume.docx'.")
    else:
        # Save the original resume if no changes are needed
        save_resume_to_file(resume_text, 'improved_resume.docx')
        print("\nYour resume meets the ATS criteria. It has been saved as 'improved_resume.docx'.")

def save_resume_to_file(resume_text, file_name):
    """
    Saves the resume text to a Word document.
    """
    document = Document()
    sections = resume_text.strip().split('\n\n')
    for section in sections:
        lines = section.strip().split('\n')
        if lines:
            # Assumes the first line is the section heading
            heading = lines[0]
            content_lines = lines[1:]
            # Add heading
            document.add_heading(heading, level=1)
            # Add content
            for line in content_lines:
                document.add_paragraph(line)
    document.save(file_name)

if __name__ == '__main__':
    main()